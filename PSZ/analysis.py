import psycopg2
import docx
from docx.shared import Pt


class DataAnalysis:

    def __init__(self, config):
        postgres_config = config['Postgres']
        self.postgres_connection = psycopg2.connect(database=postgres_config['Database'],
                                                    user=postgres_config['Username'],
                                                    password=postgres_config['Password'],
                                                    host=postgres_config['Host'],
                                                    port=postgres_config['Port'])
        self.postgres_cursor = self.postgres_connection.cursor()

        self.save_doc = docx.Document()
        self.file_name = config['FileNames']['DataAnalysis']
        self.parameters = config['DataAnalysis']['max_parameters'].split(',')

        style = self.save_doc.styles['Title']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(20)

    def num_of_cars_per_make(self):
        sql_query = '''SELECT make, count(*) FROM car GROUP BY make'''
        self.postgres_cursor.execute(sql_query)

        self.set_title("Number of cars per make:")
        for record in self.postgres_cursor:
            self.save_doc.add_paragraph(str(record))

    def num_of_cars_per_location(self):
        sql_query = '''SELECT place, count(*) FROM car GROUP BY place'''
        self.postgres_cursor.execute(sql_query)

        self.set_title("Number of cars per location:")
        for record in self.postgres_cursor:
            self.save_doc.add_paragraph(str(record))

    def num_of_cars_per_color(self):
        sql_query = '''SELECT exterior_color, count(*) FROM car GROUP BY exterior_color'''
        self.postgres_cursor.execute(sql_query)

        self.set_title("Number of cars per color:")
        for record in self.postgres_cursor:
            self.save_doc.add_paragraph(str(record))

    def most_expensive_cars(self):
        sql_query = '''SELECT * FROM car ORDER BY price DESC LIMIT 30'''
        sql_query2 = '''SELECT * FROM car WHERE body = 'Džip/SUV' ORDER BY price DESC LIMIT 30'''
        self.postgres_cursor.execute(sql_query)

        self.set_title("Top 30 cars per price:")
        for record in self.postgres_cursor:
            self.save_doc.add_paragraph(str(record))

        self.set_title("Top 30 Jeep/SUVs per price:")
        self.postgres_cursor.execute(sql_query2)
        for record in self.postgres_cursor:
            self.save_doc.add_paragraph(str(record))

    def newest_cars(self):
        sql_query = '''SELECT * FROM car WHERE year = 2021 or year = 2022 ORDER BY price 
        DESC '''
        self.postgres_cursor.execute(sql_query)

        self.set_title("Cars from this and the last year:")
        self.postgres_cursor.execute(sql_query)
        for record in self.postgres_cursor:
            self.save_doc.add_paragraph(str(record))

    def max_parameters_cars(self):
        for parameter in self.parameters:
            self.max_parameter(parameter)

    def max_parameter(self, parameter):
        sql_query = '''SELECT * FROM car where {p} = (SELECT max({p}) FROM car)'''.format(p=parameter)
        self.postgres_cursor.execute(sql_query)

        self.set_title("Cars with the biggest {p}".format(p=parameter))
        self.postgres_cursor.execute(sql_query)
        for record in self.postgres_cursor:
            self.save_doc.add_paragraph(str(record))

    def set_title(self, title_text):
        p = self.save_doc.add_paragraph()
        runner = p.add_run(title_text)
        runner.bold = True
        p.style = self.save_doc.styles['Title']


    def close(self):
        self.postgres_cursor.close()
        self.postgres_connection.commit()
        self.postgres_connection.close()
        self.save_doc.save(self.file_name)
